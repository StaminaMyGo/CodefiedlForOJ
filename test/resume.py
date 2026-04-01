from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

def create_resume():
    # 1. 创建文档
    doc = Document()
    
    # 设置页面边距
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # ========== 全局样式定义 ==========
    # 标题样式（绿色大标题）
    title_style = doc.styles.add_style('CustomTitle', 1)
    title_font = title_style.font
    title_font.name = '微软雅黑'
    title_font.size = Pt(20)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 102, 0)  # 深绿色

    # 二级标题样式（绿色小标题）
    subtitle_style = doc.styles.add_style('CustomSubtitle', 1)
    subtitle_font = subtitle_style.font
    subtitle_font.name = '微软雅黑'
    subtitle_font.size = Pt(14)
    subtitle_font.bold = True
    subtitle_font.color.rgb = RGBColor(0, 153, 0)  # 中绿色

    # 正文样式（黑色常规）
    body_style = doc.styles.add_style('CustomBody', 1)
    body_font = body_style.font
    body_font.name = '微软雅黑'
    body_font.size = Pt(11)
    body_font.color.rgb = RGBColor(0, 0, 0)

    # 列表项样式（缩进）
    list_style = doc.styles.add_style('CustomList', 1)
    list_font = list_style.font
    list_font.name = '微软雅黑'
    list_font.size = Pt(11)
    list_font.color.rgb = RGBColor(0, 0, 0)

    # ========== 简历内容填充 ==========
    # 1. 标题栏
    title_para = doc.add_paragraph(style='CustomTitle')
    title_para.add_run('个人简历').bold = True
    title_para.add_run('     ').bold = False
    title_para.add_run('Personal resume').italic = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 2. 基本信息（表格布局）
    doc.add_paragraph()  # 空行
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    info_table.autofit = True

    # 填充基本信息
    info_data = [
        ['姓 名：全民简历', '年 龄：27岁'],
        ['性 别：男', '籍 贯：上海'],
        ['工作年限：3年经验', '求职岗位：行政专员'],
        ['电 话：15688888882', '邮 箱：qmjianli@qq.com'],
        ['', '']  # 预留头像位置
    ]

    for i, row_data in enumerate(info_data):
        row_cells = info_table.rows[i].cells
        for j, cell_text in enumerate(row_data):
            cell = row_cells[j]
            cell.text = cell_text
            # 设置单元格格式
            for para in cell.paragraphs:
                para.style = body_style
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 3. 教育背景
    doc.add_paragraph('教育背景', style='CustomSubtitle')
    edu_para = doc.add_paragraph(style='CustomBody')
    edu_para.add_run('2012-09 ~ 2016-07    全民简历师范大学    工商管理（本科）').bold = True
    edu_para.add_run('\n专业成绩：GPA 3.66/4（专业前5%）').bold = False
    
    # 主修课程
    course_para = doc.add_paragraph(style='CustomList')
    course_para.add_run('主修课程：').bold = True
    course_para.add_run('基础会计学、货币银行学、统计学、经济法概论、财务会计学、管理学原理、组织行为学、市场营销学、国际贸易理论、国际贸易实务、人力资源开发与管理、财务管理学、企业经营战略概论、质量管理学、西方经济学等等。')

    # 4. 工作经验
    doc.add_paragraph('工作经验', style='CustomSubtitle')
    
    # 第一份工作
    job1_para = doc.add_paragraph(style='CustomBody')
    job1_para.add_run('2018-09 ~ 至今    全民简历科技有限公司    行政专员').bold = True
    job1_list = doc.add_paragraph(style='CustomList')
    job1_list.add_run('负责本部的行政人事管理和日常事务，协助总监搞好各部门之间的综合协调，落实公司规章制度，沟通内外联系，保证上情下达和下情上报，负责对会议文件决定的事项进行催办、查办和落实，负责全公司组织系统研讨和修订。').bold = False
    
    job1_list2 = doc.add_paragraph(style='CustomList')
    job1_list2.add_run('编制公司人事管理制度，规避各项人事风险。').bold = False

    # 第二份工作
    job2_para = doc.add_paragraph(style='CustomBody')
    job2_para.add_run('2016-09 ~ 2018-08    上海斧掌网络科技有限公司    行政专员').bold = True
    job2_list1 = doc.add_paragraph(style='CustomList')
    job2_list1.add_run('负责中心简单财务管理，资产管控；').bold = False
    
    job2_list2 = doc.add_paragraph(style='CustomList')
    job2_list2.add_run('负责公司总部的来访客户接待工作，负责引导和介绍公司的分布情况；').bold = False
    
    job2_list3 = doc.add_paragraph(style='CustomList')
    job2_list3.add_run('负责中心的行政事务，公司班车管理、负责建立员工归属感及前台管理；').bold = False
    
    job2_list4 = doc.add_paragraph(style='CustomList')
    job2_list4.add_run('负责招聘工作，确保人才梯队发展和人才储备及培养；').bold = False
    
    job2_list5 = doc.add_paragraph(style='CustomList')
    job2_list5.add_run('督导公司各项行政、人事制度、员工福利、生日以及公司各种宴会活动的执行；').bold = False
    
    job2_list6 = doc.add_paragraph(style='CustomList')
    job2_list6.add_run('负责招聘工作，制定公司的人力资源发展计划，确保人才梯队发展和人才储备及培养。').bold = False

    # 5. 技能特长
    doc.add_paragraph('技能特长', style='CustomSubtitle')
    skill_para1 = doc.add_paragraph(style='CustomBody')
    skill_para1.add_run('语言能力：').bold = True
    skill_para1.add_run('大学英语6级证书，荣获全国大学生英语竞赛一等奖，能够熟练的进行交流、读写。')
    
    skill_para2 = doc.add_paragraph(style='CustomBody')
    skill_para2.add_run('计算机：').bold = True
    skill_para2.add_run('计算机二级证书，熟练操作windows平台上的各类应用软件，如Word、Excel。')
    
    skill_para3 = doc.add_paragraph(style='CustomBody')
    skill_para3.add_run('团队能力：').bold = True
    skill_para3.add_run('具有丰富的团队组建与扩充经验和项目管理与协调经验。')

    # 6. 荣誉证书
    doc.add_paragraph('荣誉证书', style='CustomSubtitle')
    honor_list1 = doc.add_paragraph(style='CustomList')
    honor_list1.add_run('英语四级，听说读写能力良好，能流利的用英语进行日常交流，能快速浏览英文文档和书籍；').bold = False
    
    honor_list2 = doc.add_paragraph(style='CustomList')
    honor_list2.add_run('通过全国计算机二级考试，熟练运用office等常用的办公软件。').bold = False

    # 7. 自我评价
    doc.add_paragraph('自我评价', style='CustomSubtitle')
    self_para = doc.add_paragraph(style='CustomBody')
    self_para.add_run('工作积极认真，细心负责，熟练运用办公自动化软件，善于在工作中提出问题、发现问题、解决问题，有较强的分析能力；勤奋好学，踏实肯干，动手能力强，认真负责，有很强的社会责任感；坚毅不拔，吃苦耐劳，喜欢迎接新挑战。').bold = False

    # 8. 保存文档
    doc.save('定制化绿色风简历.docx')
    print("简历生成成功：定制化绿色风简历.docx")

if __name__ == "__main__":
    create_resume()